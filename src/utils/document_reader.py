#!/usr/bin/env python3
"""
Módulo para leer y procesar documentos (PDF, Excel, etc.)
"""

import os
import logging
from typing import Dict, List, Any, Optional
import pandas as pd
import PyPDF2
import pdfplumber
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.colors import black
import io
from pdf2image import convert_from_path
from PIL import Image, ImageDraw, ImageFont
import fitz  # PyMuPDF
from datetime import datetime

logger = logging.getLogger(__name__)

class DocumentReader:
    """Clase para leer y procesar diferentes tipos de documentos"""
    
    def __init__(self):
        """Inicializa el lector de documentos"""
        self.supported_extensions = {
            '.pdf': self._read_pdf,
            '.xlsx': self._read_excel,
            '.xls': self._read_excel,
            '.docx': self._read_docx,
            '.doc': self._read_docx
        }
    
    def read_document(self, file_path: str) -> Dict[str, Any]:
        """
        Lee un documento y devuelve su contenido procesado
        
        Args:
            file_path: Ruta al archivo a leer
            
        Returns:
            Diccionario con el contenido procesado
        """
        if not os.path.exists(file_path):
            logger.error(f"Archivo no encontrado: {file_path}")
            return {}
            
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext not in self.supported_extensions:
            logger.error(f"Extensión no soportada: {file_ext}")
            return {}
            
        try:
            return self.supported_extensions[file_ext](file_path)
        except Exception as e:
            logger.error(f"Error leyendo archivo {file_path}: {str(e)}")
            return {}
    
    def _read_pdf(self, file_path: str) -> Dict[str, Any]:
        """Lee un archivo PDF usando múltiples métodos"""
        content = {
            'text': '',
            'metadata': {},
            'pages': []
        }
        
        # Intentar con PyPDF2 primero
        try:
            with open(file_path, 'rb') as file:
                pdf = PyPDF2.PdfReader(file)
                
                # Obtener metadatos
                if pdf.metadata:
                    content['metadata'] = {
                        'title': pdf.metadata.get('/Title', ''),
                        'author': pdf.metadata.get('/Author', ''),
                        'subject': pdf.metadata.get('/Subject', ''),
                        'creator': pdf.metadata.get('/Creator', '')
                    }
                
                # Extraer texto de cada página
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content['pages'].append(text)
                        content['text'] += text + '\n\n'
                        
        except Exception as e:
            logger.warning(f"Error con PyPDF2, intentando pdfplumber: {str(e)}")
            
            # Fallback a pdfplumber
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        text = page.extract_text()
                        if text:
                            content['pages'].append(text)
                            content['text'] += text + '\n\n'
            except Exception as e:
                logger.error(f"Error con pdfplumber: {str(e)}")
        
        return content
    
    def _read_excel(self, file_path: str) -> Dict[str, Any]:
        """Lee un archivo Excel"""
        content = {
            'sheets': {},
            'metadata': {}
        }
        
        try:
            # Leer todas las hojas
            excel_file = pd.ExcelFile(file_path)
            
            # Obtener metadatos
            content['metadata'] = {
                'sheets': excel_file.sheet_names,
                'file_name': os.path.basename(file_path)
            }
            
            # Leer cada hoja
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convertir DataFrame a diccionario
                sheet_data = {
                    'headers': df.columns.tolist(),
                    'data': df.fillna('').to_dict('records')
                }
                
                content['sheets'][sheet_name] = sheet_data
                
        except Exception as e:
            logger.error(f"Error leyendo Excel {file_path}: {str(e)}")
        
        return content
    
    def _read_docx(self, file_path: str) -> Dict[str, Any]:
        """Lee un archivo Word"""
        content = {
            'text': '',
            'metadata': {},
            'paragraphs': []
        }
        
        try:
            doc = Document(file_path)
            
            # Obtener metadatos
            content['metadata'] = {
                'title': doc.core_properties.title,
                'author': doc.core_properties.author,
                'created': str(doc.core_properties.created),
                'modified': str(doc.core_properties.modified)
            }
            
            # Extraer texto de cada párrafo
            for para in doc.paragraphs:
                if para.text.strip():
                    content['paragraphs'].append(para.text)
                    content['text'] += para.text + '\n'
                    
        except Exception as e:
            logger.error(f"Error leyendo Word {file_path}: {str(e)}")
        
        return content
    
    def print_document_content(self, file_path: str) -> None:
        """
        Imprime el contenido de un documento en la terminal
        
        Args:
            file_path: Ruta al archivo a leer
        """
        content = self.read_document(file_path)
        
        if not content:
            print(f"❌ No se pudo leer el archivo: {file_path}")
            return
            
        file_ext = os.path.splitext(file_path)[1].lower()
        file_name = os.path.basename(file_path)
        
        print(f"\n📄 Contenido de {file_name}:")
        print("=" * 80)
        
        if file_ext in ['.xlsx', '.xls']:
            for sheet_name, sheet_data in content['sheets'].items():
                print(f"\n📊 Hoja: {sheet_name}")
                print("-" * 40)
                
                # Imprimir encabezados
                headers = sheet_data['headers']
                print(" | ".join(headers))
                print("-" * 40)
                
                # Imprimir datos
                for row in sheet_data['data']:
                    print(" | ".join(str(row[h]) for h in headers))
                    
        elif file_ext in ['.pdf', '.docx', '.doc']:
            if content.get('metadata'):
                print("\n📋 Metadatos:")
                for key, value in content['metadata'].items():
                    if value:
                        print(f"{key}: {value}")
                print("-" * 40)
            
            if content.get('text'):
                print("\n📝 Contenido:")
                print(content['text'])
                
        print("\n" + "=" * 80)
    
    def validate_excel_structure(self, file_path: str, reference_path: str) -> bool:
        """
        Valida si un archivo Excel tiene la misma estructura (encabezados) que una plantilla de referencia
        
        Args:
            file_path: Ruta al archivo a validar
            reference_path: Ruta al archivo de referencia
        Returns:
            True si la estructura es igual, False si no
        """
        try:
            excel_file = pd.ExcelFile(file_path)
            ref_file = pd.ExcelFile(reference_path)
            
            # Comprobar número de hojas
            if len(excel_file.sheet_names) != len(ref_file.sheet_names):
                return False
            
            for sheet1, sheet2 in zip(excel_file.sheet_names, ref_file.sheet_names):
                df1 = pd.read_excel(file_path, sheet_name=sheet1)
                df2 = pd.read_excel(reference_path, sheet_name=sheet2)
                
                if list(df1.columns) != list(df2.columns):
                    return False
            return True
        except Exception as e:
            logger.error(f"Error validando estructura de Excel: {str(e)}")
            return False
    
    def customize_application_form(self, template_path: str, output_path: str, offer_data: Dict) -> Optional[str]:
        """
        Personaliza archivos .docx (Application Form) con los datos de la oferta.
        Solo modifica los campos POSITION ADVERTISED, School y ROLL NUMBER sin alterar el resto del documento.
        
        Args:
            template_path: Ruta a la plantilla del documento (.docx)
            output_path: Ruta donde guardar el archivo personalizado
            offer_data: Diccionario con los datos de la oferta (position, school_name, roll_number)
            
        Returns:
            Ruta al archivo personalizado o None si hay un error
        """
        if not os.path.exists(template_path) or not template_path.lower().endswith('.docx'):
            logger.error(f"Plantilla no encontrada o formato incorrecto: {template_path}")
            return None
            
        try:
            # Asegurar que existe el directorio de salida
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Verificar que tenemos los datos necesarios
            position = offer_data.get('position', 'Teaching Position')
            school_name = offer_data.get('school_name', 'School')
            
            # Limpiar el roll number para eliminar cualquier sufijo como "Apply"
            roll_number = offer_data.get('roll_number', 'N/A')
            if roll_number and isinstance(roll_number, str):
                # Extraer solo la parte numérica o código del roll number
                if "Apply" in roll_number:
                    clean_roll_number = roll_number.split("Apply")[0].strip()
                    logger.info(f"Roll Number limpiado: '{roll_number}' -> '{clean_roll_number}'")
                    roll_number = clean_roll_number
            
            logger.info(f"Personalizando application form con:")
            logger.info(f"- Position: {position}")
            logger.info(f"- School: {school_name}")
            logger.info(f"- Roll Number: {roll_number}")
            
            # Cargar el documento
            doc = Document(template_path)
            
            # Variables para seguimiento de reemplazos
            position_replaced = False
            school_replaced = False
            roll_number_replaced = False
            
            # Función auxiliar para reemplazar en un párrafo preservando formato
            def replace_in_paragraph(paragraph):
                nonlocal position_replaced, school_replaced, roll_number_replaced
                
                # Usamos runs para preservar el formato de cada parte del texto
                for i, run in enumerate(paragraph.runs):
                    # Reemplazar POSITION ADVERTISED
                    if 'POSITION ADVERTISED' in run.text:
                        run.text = run.text.replace('POSITION ADVERTISED', position)
                        position_replaced = True
                        logger.info(f"Reemplazado 'POSITION ADVERTISED' con '{position}' en run {i}")
                
                    # Reemplazar School: [value]
                    if 'School:' in run.text:
                        parts = run.text.split('School:', 1)
                        if len(parts) > 1:
                            # Preservar el texto antes de "School:"
                            new_text = parts[0] + f"School: {school_name}"
                            # Si hay texto después del valor de la escuela, preservarlo
                            if ':' in parts[1] and len(parts[1].split(':', 1)) > 1:
                                next_part = parts[1].split(':', 1)[1]
                                new_text = parts[0] + f"School: {school_name}:" + next_part
                            run.text = new_text
                            school_replaced = True
                            logger.info(f"Reemplazado texto después de 'School:' con '{school_name}' en run {i}")
                
                    # Reemplazar ROLL NUMBER
                    if 'ROLL NUMBER' in run.text:
                        run.text = run.text.replace('ROLL NUMBER', roll_number)
                        roll_number_replaced = True
                        logger.info(f"Reemplazado 'ROLL NUMBER' con '{roll_number}' en run {i}")
            
            # Recorrer todos los párrafos buscando los campos a reemplazar
            for paragraph in doc.paragraphs:
                replace_in_paragraph(paragraph)
            
            # También buscar en las tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            replace_in_paragraph(paragraph)
            
            # Log de resultados
            logger.info(f"Resultados de los reemplazos:")
            logger.info(f"Position replaced: {position_replaced}")
            logger.info(f"School replaced: {school_replaced}")
            logger.info(f"Roll number replaced: {roll_number_replaced}")
            
            # Guardar el documento modificado
            doc.save(output_path)
            logger.info(f"Application Form personalizado guardado en: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Error personalizando Application Form: {str(e)}")
            return None
    
    def customize_application_form_pdf(self, template_path: str, output_path: str, offer_data: Dict) -> Optional[str]:
        """
        Personaliza archivos PDF (Application Form) con los datos de la oferta.
        Ahora la primera página se renderiza como imagen y se personaliza visualmente.
        El resto de páginas se copian tal cual.
        """
        if not os.path.exists(template_path) or not template_path.lower().endswith('.pdf'):
            logger.error(f"Plantilla PDF no encontrada o formato incorrecto: {template_path}")
            return None
        try:
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            position = offer_data.get('position', 'Teaching Position')
            school_name = offer_data.get('school_name', 'School')
            roll_number = offer_data.get('roll_number', 'N/A')
            if roll_number and isinstance(roll_number, str):
                if "Apply" in roll_number:
                    clean_roll_number = roll_number.split("Apply")[0].strip()
                    roll_number = clean_roll_number

            campos = {
                'POSITION ADVERTISED': f"POSITION ADVERTISED : {position}",
                'School:': f"School: {school_name}",
                'ROLL NUMBER': f"ROLL NUMBER: {roll_number}"
            }
            doc = fitz.open(template_path)
            page = doc[0]
            # Variantes de los campos a buscar
            variantes = {
                'POSITION ADVERTISED': [
                    'POSITION ADVERTISED',
                    'POSITION ADVERTISED:',
                    'Position Advertised',
                    'Position Advertised:',
                ],
                'School:': [
                    'School:',
                    'SCHOOL:',
                    'School',
                    'SCHOOL'
                ],
                'ROLL NUMBER': [
                    'ROLL NUMBER',
                    'ROLL NUMBER:',
                    'Roll Number',
                    'Roll Number:',
                ]
            }
            font_size = 12  # Tamaño de fuente fijo y profesional
            roll_rect = None
            # 1. Sobrescribir ROLL NUMBER (como hasta ahora)
            key = 'ROLL NUMBER'
            new_text = campos[key]
            topmost_rect = None
            topmost_y = float('inf')
            for variante in variantes[key]:
                areas = page.search_for(variante, quads=False)
                for rect in areas:
                    if rect.y0 < topmost_y:
                        topmost_y = rect.y0
                        topmost_rect = rect
            if topmost_rect:
                roll_rect = topmost_rect
                padding = 2
                rect_expanded = fitz.Rect(topmost_rect.x0 - padding, topmost_rect.y0 - padding, topmost_rect.x1 + 120, topmost_rect.y1 + padding)
                page.draw_rect(rect_expanded, color=(1,1,1), fill=(1,1,1))
                page.insert_text((topmost_rect.x0, topmost_rect.y0), new_text, fontsize=font_size, color=(0,0,0))

            # 2. Sobrescribir School en la coincidencia más cercana a ROLL NUMBER
            key = 'School:'
            new_text = campos[key]
            closest_rect = None
            min_distance = float('inf')
            for variante in variantes[key]:
                areas = page.search_for(variante, quads=False)
                for rect in areas:
                    if roll_rect:
                        distance = abs(rect.y0 - roll_rect.y0)
                        if distance < min_distance:
                            min_distance = distance
                            closest_rect = rect
            if closest_rect:
                padding = 2
                rect_expanded = fitz.Rect(closest_rect.x0 - padding, closest_rect.y0 - padding, closest_rect.x1 + 120, closest_rect.y1 + padding)
                page.draw_rect(rect_expanded, color=(1,1,1), fill=(1,1,1))
                page.insert_text((closest_rect.x0, closest_rect.y0), new_text, fontsize=font_size, color=(0,0,0))

            # 3. Sobrescribir POSITION ADVERTISED (como antes, la más arriba)
            key = 'POSITION ADVERTISED'
            new_text = campos[key]
            topmost_rect = None
            topmost_y = float('inf')
            for variante in variantes[key]:
                areas = page.search_for(variante, quads=False)
                for rect in areas:
                    if rect.y0 < topmost_y:
                        topmost_y = rect.y0
                        topmost_rect = rect
            if topmost_rect:
                padding = 2
                # Ampliar mucho más el rectángulo blanco para tapar cualquier resto de vacante anterior
                rect_expanded = fitz.Rect(topmost_rect.x0 - padding, topmost_rect.y0 - padding, topmost_rect.x0 + 350, topmost_rect.y1 + padding)
                page.draw_rect(rect_expanded, color=(1,1,1), fill=(1,1,1))
                page.insert_text((topmost_rect.x0, topmost_rect.y0), new_text, fontsize=font_size, color=(0,0,0))
            
            # 4. Añadir la fecha en la última página
            try:
                last_page = doc[-1]
                date_variantes = ['Date:', 'Date', 'DATE:', 'DATE']
                date_rect = None
                
                # Buscar el texto "Date:" en la última página
                for variante in date_variantes:
                    areas = last_page.search_for(variante, quads=False)
                    if areas:
                        # Usar la primera coincidencia que encontremos
                        date_rect = areas[0]
                        break
                
                if date_rect:
                    current_date = datetime.now().strftime("%d/%m/%Y")
                    # La posición de inserción es a la derecha de la etiqueta "Date:"
                    # Ajuste: se reduce el espacio a 5px y se sube 2px para mejor alineación
                    insert_pos = fitz.Point(date_rect.x1 + 5, date_rect.y1 - 2)
                    
                    # Definir el rectángulo para tapar la fecha anterior
                    # Se extiende desde la derecha de "Date:" hasta el final del área esperada de la fecha
                    cover_rect = fitz.Rect(
                        date_rect.x1 + 1,   # Empezar a cubrir 1px después de "Date:"
                        date_rect.y0 - 2, # Un poco de padding vertical
                        date_rect.x1 + 150, # Ancho suficiente para tapar fechas largas
                        date_rect.y1 + 2
                    )
                    
                    # Dibujar el rectángulo blanco para tapar la fecha vieja
                    last_page.draw_rect(cover_rect, color=(1, 1, 1), fill=(1, 1, 1))
                    
                    # Insertar el texto de la fecha nueva
                    last_page.insert_text(insert_pos, current_date, fontsize=font_size, color=(0,0,0))
                    logger.info(f"Fecha '{current_date}' insertada en la última página (cubriendo la anterior).")
                else:
                    logger.warning("No se encontró el campo 'Date:' en la última página. No se pudo insertar la fecha.")
            except Exception as e:
                logger.error(f"Error al intentar añadir la fecha en la última página: {str(e)}")

            doc.save(output_path)
            doc.close()
            logger.info(f"Application Form PDF personalizado guardado en: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Error personalizando Application Form PDF: {str(e)}")
            return None